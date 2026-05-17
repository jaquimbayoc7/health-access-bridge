"""
Generador del documento HCI - Actividad 1: Contextualización del Proyecto HAB
Ejecutar: py docs/generate_hci_doc.py
Salida:   docs/HCI_Actividad1_HAB.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMAGES_DIR = os.path.join(BASE_DIR, "..", "frontend", "design", "images")
OUTPUT_PATH = os.path.join(BASE_DIR, "HCI_Actividad1_HAB.docx")


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color="1F3864"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_paragraph(doc, text, bold=False, italic=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_image_with_caption(doc, img_path, caption, width=Inches(5.5)):
    if not os.path.exists(img_path):
        doc.add_paragraph(f"[Imagen no encontrada: {os.path.basename(img_path)}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=width)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].italic = True
    cap.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    cap.paragraph_format.space_after = Pt(14)


def build_document():
    doc = Document()

    # --- Márgenes ---
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # =========================================================
    # PORTADA
    # =========================================================
    portada_path = os.path.join(IMAGES_DIR, "Portada.png")
    if os.path.exists(portada_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(portada_path, width=Inches(5.8))

    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HCI — Actividad 1\nContextualización del Proyecto")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Health Access Bridge (HAB)")
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    run2.bold = True

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = meta.add_run("Asignatura: Interacción Humano-Computador (HCI)\nPosgrado · Proyecto Integrador\nAbril 2026")
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_page_break()

    # =========================================================
    # SECCIÓN 1 — PROBLEMA
    # =========================================================
    add_heading(doc, "1. Descripción del Problema", level=1)

    add_heading(doc, "1.1 Contexto y magnitud", level=2, color="2E74B5")
    add_paragraph(doc,
        "En Colombia y Latinoamérica, las personas con discapacidad que habitan en zonas rurales "
        "enfrentan barreras sistemáticas para acceder a los servicios de salud. Estas barreras son "
        "de naturaleza múltiple: distancia geográfica y ausencia de transporte adaptado, falta de "
        "personal médico especializado en territorios alejados, brecha digital que impide el acceso "
        "a servicios en línea, barreras económicas que elevan el costo efectivo de la atención, y "
        "barreras culturales y de comunicación que dificultan la relación médico-paciente."
    )
    add_paragraph(doc,
        "Según el DANE, más del 7 % de la población colombiana presenta algún tipo de discapacidad, "
        "y la prevalencia es significativamente mayor en zonas rurales dispersas. Sin embargo, los "
        "sistemas de información clínica disponibles no están diseñados para capturar, estructurar ni "
        "analizar el perfil de barreras de acceso de esta población."
    )

    add_heading(doc, "1.2 Dolores principales de los actores involucrados", level=2, color="2E74B5")

    pain_points = [
        ("Médico rural / de familia",
         "No cuenta con herramientas digitales integradas para registrar el perfil funcional de sus "
         "pacientes con discapacidad ni para anticipar sus barreras de acceso. La gestión es manual, "
         "fragmentada y sin soporte de inteligencia artificial."),
        ("Administrador del sistema de salud",
         "Carece de visibilidad poblacional en tiempo real. No puede exportar reportes ni tomar "
         "decisiones basadas en datos agregados sobre distribución de perfiles de discapacidad."),
        ("Sistema de salud (macro)",
         "Ausencia de datos estructurados impide la formulación de políticas de intervención "
         "diferencial para poblaciones rurales con discapacidad.")
    ]

    for actor, desc in pain_points:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{actor}: ")
        run.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(desc)
        run2.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    # =========================================================
    # SECCIÓN 2 — IDEA DE SOLUCIÓN
    # =========================================================
    add_heading(doc, "2. Idea de Solución", level=1)

    add_heading(doc, "2.1 Descripción de HAB", level=2, color="2E74B5")
    add_paragraph(doc,
        "Health Access Bridge (HAB) es una plataforma web inteligente orientada a la gestión clínica "
        "de pacientes con discapacidad y al análisis predictivo de sus perfiles de barreras de acceso "
        "a la salud en zonas rurales. La solución conecta a médicos rurales con herramientas digitales "
        "de registro, análisis y apoyo a la decisión clínica, reduciendo la brecha tecnológica en "
        "entornos de baja conectividad."
    )

    add_heading(doc, "2.2 Componente de Inteligencia Artificial", level=2, color="2E74B5")
    add_paragraph(doc,
        "HAB integra dos capas de inteligencia artificial:"
    )
    ai_items = [
        ("Modelo predictivo ML (HybridModelDisability)",
         "Modelo entrenado con K-Means + Gradient Boosting que clasifica automáticamente a cada "
         "paciente en uno de tres perfiles de barreras de acceso (Perfil 0, 1 ó 2), "
         "basándose en sus datos clínicos y sociodemográficos. El modelo está embebido en el backend "
         "y se invoca mediante el endpoint POST /patients/{id}/predict."),
        ("LLM para recomendaciones clínicas (roadmap)",
         "Integración futura con GPT-4 u otro LLM para generar resúmenes clínicos personalizados "
         "y planes de intervención basados en el perfil ICF del paciente, con opción de exportación "
         "a PDF y disclaimer legal.")
    ]
    for title_ai, desc_ai in ai_items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{title_ai}: ")
        run.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(desc_ai)
        run2.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    # =========================================================
    # SECCIÓN 3 — FUNCIONALIDADES POTENCIALES
    # =========================================================
    add_heading(doc, "3. Funcionalidades Potenciales", level=1)
    add_paragraph(doc,
        "A continuación se listan las funcionalidades identificadas para la solución HAB:"
    )

    funcionalidades = [
        ("F1",  "Autenticación con roles diferenciados (Administrador / Médico) — RBAC"),
        ("F2",  "Registro, edición y gestión de historia clínica de pacientes con discapacidad"),
        ("F3",  "Búsqueda, filtrado y paginación de pacientes"),
        ("F4",  "Predicción automática del perfil de barreras de acceso (modelo ML embebido)"),
        ("F5",  "Guía clínica interpretativa de los 3 perfiles predictivos"),
        ("F6",  "Dashboard de analíticas con distribución de perfiles (gráfica de torta)"),
        ("F7",  "Exportación de datos a Excel / PDF"),
        ("F8",  "Modo offline / PWA para zonas de baja conectividad"),
        ("F9",  "Recomendaciones clínicas generadas por LLM (GPT-4)"),
        ("F10", "Panel de administración: gestión de usuarios, activación/desactivación"),
        ("F11", "Despliegue multi-ambiente (DEV / QA / PROD) con CI/CD automatizado"),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Encabezado
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Funcionalidad"
    for i, cell in enumerate(hdr):
        set_cell_bg(cell, "1F3864")
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(11)

    # Filas
    for idx, (code, desc) in enumerate(funcionalidades):
        row = table.add_row().cells
        row[0].text = code
        row[1].text = desc
        bg = "D6E4F0" if idx % 2 == 0 else "FFFFFF"
        set_cell_bg(row[0], bg)
        set_cell_bg(row[1], bg)
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10.5)

    # Ancho columnas
    table.columns[0].width = Cm(1.5)
    table.columns[1].width = Cm(13)

    doc.add_paragraph()

    # =========================================================
    # SECCIÓN 4 — PROTOTIPOS PRELIMINARES
    # =========================================================
    add_heading(doc, "4. Prototipos Preliminares — Mockups de Pantallas", level=1)
    add_paragraph(doc,
        "Las siguientes imágenes corresponden a los mockups de diseño preliminares de las "
        "8 pantallas principales de la aplicación HAB, desarrollados como artefacto de la "
        "Épica 1 del proyecto."
    )

    screens = [
        ("1-Login.png",          "Figura 1. Pantalla de Login — Autenticación por roles (Admin / Médico)"),
        ("2-Admin.png",          "Figura 2. Panel Administrador — Gestión de usuarios y visión global"),
        ("3-DashMedico.png",     "Figura 3. Dashboard Médico — Resumen de pacientes y métricas clave"),
        ("4-Pacientes.png",      "Figura 4. Gestión de Pacientes — Lista, búsqueda, edición y eliminación"),
        ("5-Predicciones.png",   "Figura 5. Módulo de Predicciones — Ejecución del modelo ML e historial"),
        ("6-Analisis.png",       "Figura 6. Análisis — Dashboard con distribución de perfiles predictivos"),
        ("7-GuiaPrediccion.png", "Figura 7. Guía de Predicción — Interpretación clínica de los 3 perfiles"),
        ("8-PerfilICF.png",      "Figura 8. Perfil Funcional ICF — Clasificación Internacional del Funcionamiento"),
    ]

    for filename, caption in screens:
        img_path = os.path.join(IMAGES_DIR, filename)
        add_image_with_caption(doc, img_path, caption)

    doc.add_paragraph()

    # =========================================================
    # SECCIÓN 5 — USUARIOS FINALES
    # =========================================================
    add_heading(doc, "5. Identificación de Usuarios Finales Potenciales", level=1)
    add_paragraph(doc,
        "A continuación se identifican los segmentos de usuarios finales potenciales para la "
        "solución HAB, incluyendo una descripción de su perfil y sus necesidades clave:"
    )

    usuarios = [
        ("Médico rural / de familia",
         "Profesional de salud que atiende pacientes con discapacidad en zonas de difícil acceso. "
         "Maneja múltiples casos con recursos limitados y necesita herramientas ágiles.",
         "Registro rápido de historia clínica, predicción automatizada de perfiles, acceso offline, "
         "historial clínico centralizado"),
        ("Administrador del sistema de salud",
         "Coordinador o directivo que supervisa equipos médicos y requiere visión global de la "
         "población atendida para tomar decisiones estratégicas.",
         "Dashboard de analíticas, gestión de cuentas de usuarios, exportación de reportes a PDF/Excel"),
        ("Especialista en rehabilitación / trabajo social",
         "Profesional de apoyo que consulta el perfil funcional del paciente para diseñar planes "
         "de intervención individualizados basados en el marco ICF.",
         "Perfil Funcional ICF, guía clínica de barreras, recomendaciones generadas por LLM (futuro)"),
        ("Paciente con discapacidad (usuario indirecto)",
         "No interactúa directamente con el sistema, pero es el beneficiario final. Su información "
         "clínica y perfil de barreras orienta las decisiones médicas.",
         "Calidad y exactitud del dato registrado, confidencialidad y privacidad de la información"),
        ("Investigador / epidemiólogo",
         "Utiliza los datos agregados para estudiar tendencias y patrones de acceso a la salud en "
         "poblaciones rurales con discapacidad.",
         "Exportación masiva de datos, analíticas poblacionales, trazabilidad de registros históricos"),
    ]

    table2 = doc.add_table(rows=1, cols=3)
    table2.style = "Table Grid"
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers2 = ["Segmento de usuario", "Descripción", "Necesidades clave"]
    hdr2 = table2.rows[0].cells
    for i, h in enumerate(headers2):
        hdr2[i].text = h
        set_cell_bg(hdr2[i], "1F3864")
        for para in hdr2[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10.5)

    for idx, (seg, desc, needs) in enumerate(usuarios):
        row = table2.add_row().cells
        row[0].text = seg
        row[1].text = desc
        row[2].text = needs
        bg = "D6E4F0" if idx % 2 == 0 else "FFFFFF"
        for cell in row:
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    # Ancho columnas
    table2.columns[0].width = Cm(4)
    table2.columns[1].width = Cm(6)
    table2.columns[2].width = Cm(5)

    doc.add_paragraph()

    # =========================================================
    # GUARDAR
    # =========================================================
    doc.save(OUTPUT_PATH)
    print(f"Documento generado exitosamente:\n  {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
